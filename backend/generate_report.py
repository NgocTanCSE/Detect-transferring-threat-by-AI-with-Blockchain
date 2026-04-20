from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('BLOCKCHAIN SENTINEL - BÁO CÁO GIẢ LẬP RỦI RO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_heading('1. Giới thiệu', level=1)
    doc.add_paragraph(
        "Báo cáo này tổng hợp kết quả của các thử nghiệm giả lập rủi ro được thực hiện nhằm đánh giá "
        "khả năng phát hiện và ngăn chặn các hành vi gian lận trên mạng lưới blockchain của hệ thống Blockchain Sentinel."
    )

    # Scenarios
    doc.add_heading('2. Các kịch bản giả lập & Bằng chứng', level=1)

    # Money Laundering
    doc.add_heading('2.1. Rửa tiền (Smurfing/Structuring)', level=2)
    doc.add_paragraph(
        "Mô phỏng kỹ thuật chia nhỏ dòng tiền từ một ví nguồn 'bẩn' sang nhiều ví trung gian (Mules). "
        "Hệ thống theo diện rộng đã phát hiện các giao dịch lặp lại với giá trị cao (9.5 ETH)."
    )
    img_ml = r'C:\Users\Ngoc Tan\.gemini\antigravity\brain\00dfb246-6da9-4fbd-b090-aa5248734aae\money_laundering_9_5_eth_evidence_1776698381069.png'
    if os.path.exists(img_ml):
        doc.add_picture(img_ml, width=Inches(5.5))
        doc.add_paragraph("Hình 1: Các giao dịch Smurfing (9.5 ETH) được ghi nhận trong lịch sử ví nguồn.", style='Caption')

    # Spam Registration
    doc.add_heading('2.2. Tấn công Spam Đăng ký', level=2)
    doc.add_paragraph(
        "Giả lập hành vi bot đăng ký hàng loạt tài khoản với tên người dùng và email nghi vấn."
    )
    doc.add_paragraph(
        "Kết quả: Hệ thống chặn đứng 10 tài khoản bot liên tiếp, bảo vệ tài nguyên hệ thống."
    )

    # Phishing
    doc.add_heading('2.3. Tấn công Phishing (Wallet Drainer)', level=2)
    doc.add_paragraph(
        "Mô phỏng hành vi rút tiền hàng loạt từ nhiều nạn nhân về ví Scammer."
    )
    img_scam = r'C:\Users\Ngoc Tan\.gemini\antigravity\brain\00dfb246-6da9-4fbd-b090-aa5248734aae\scam_wallet_details_1776696092558.png'
    if os.path.exists(img_scam):
        doc.add_picture(img_scam, width=Inches(5.5))
        doc.add_paragraph("Hình 2: Ví Scammer bị phát hiện và gắn nhãn 'Frozen' với Risk Score 100.", style='Caption')

    # Statistics
    doc.add_heading('3. Thống kê tổng hợp vận hành', level=1)
    doc.add_paragraph(
        "Dưới đây là bảng tổng hợp các chỉ số bảo mật ghi nhận được sau khi thực hiện các bài test."
    )
    img_stats = r'C:\Users\Ngoc Tan\.gemini\antigravity\brain\00dfb246-6da9-4fbd-b090-aa5248734aae\compliance_blocked_stats_1776698226999.png'
    if os.path.exists(img_stats):
        doc.add_picture(img_stats, width=Inches(5.5))
        doc.add_paragraph("Hình 3: Tổng giá trị giao dịch rủi ro bị chặn (146 ETH) và số lượng Alert.", style='Caption')

    # Conclusion
    doc.add_heading('4. Kết luận', level=1)
    doc.add_paragraph(
        "Hệ thống đã hoạt động đúng như kỳ vọng, phát hiện chính xác các hành vi rốt ráo "
        "và áp dụng các biện pháp ngăn chặn kịp thời đối với cả 3 kịch bản tấn công rủi ro."
    )

    # Save
    output_path = r'c:\Users\Ngoc Tan\Downloads\blockchain-ai-project\Bao_Cao_Gia_Lap_Rui_Ro_Full.docx'
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

if __name__ == "__main__":
    create_report()
